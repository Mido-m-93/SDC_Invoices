"""
Generate the bilingual SDC Invoice Tool User Guide as a Word document.
Run: python generate_guide.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colours ───────────────────────────────────────────────────────────────────
GREEN_DARK = RGBColor(0x1a, 0x3d, 0x2b)
GREEN_MID  = RGBColor(0x2d, 0x6a, 0x4f)
AMBER      = RGBColor(0xd9, 0x7b, 0x06)
STONE      = RGBColor(0x78, 0x71, 0x6c)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEAD = RGBColor(0xf5, 0xf5, 0xf4)
STEP_BG    = RGBColor(0xe8, 0xf5, 0xee)


def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def section_title(doc, ja, en):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(ja)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = GREEN_DARK
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(10)
    r2 = p2.add_run(en)
    r2.font.size = Pt(9)
    r2.font.color.rgb = STONE
    r2.italic = True


def step_block(doc, number, ja_title, en_title, bullets_ja, bullets_en):
    """A numbered step with bullet sub-items."""
    # Step header row (number + title in a shaded table)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.allow_autofit = True

    num_cell = tbl.rows[0].cells[0]
    set_cell_bg(num_cell, GREEN_MID)
    num_cell.width = Cm(1.4)
    np = num_cell.paragraphs[0]
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = np.add_run(str(number))
    nr.bold = True
    nr.font.size = Pt(18)
    nr.font.color.rgb = WHITE

    title_cell = tbl.rows[0].cells[1]
    set_cell_bg(title_cell, STEP_BG)
    tp = title_cell.paragraphs[0]
    tr = tp.add_run(ja_title + "\n")
    tr.bold = True
    tr.font.size = Pt(12)
    tr.font.color.rgb = GREEN_DARK
    tr2 = tp.add_run(en_title)
    tr2.font.size = Pt(9)
    tr2.font.color.rgb = STONE
    tr2.italic = True

    doc.add_paragraph()  # small gap

    # Bullet items
    for ja, en in zip(bullets_ja, bullets_en):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(ja)
        r.font.size = Pt(10.5)

        p2 = doc.add_paragraph(style="List Bullet")
        p2.paragraph_format.left_indent  = Cm(0.8)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(6)
        r2 = p2.add_run(en)
        r2.font.size = Pt(9)
        r2.font.color.rgb = STONE
        r2.italic = True

    doc.add_paragraph()  # gap after step


def ref_table(doc, headers_ja, headers_en, rows_ja, rows_en):
    col_count = len(headers_ja)
    table = doc.add_table(rows=1 + len(rows_ja), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = table.rows[0]
    for i, (ja, en) in enumerate(zip(headers_ja, headers_en)):
        cell = hdr.cells[i]
        set_cell_bg(cell, GREEN_DARK)
        p = cell.paragraphs[0]
        r = p.add_run(f"{ja}  ")
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
        r2 = p.add_run(f"/ {en}")
        r2.font.size = Pt(7.5)
        r2.font.color.rgb = RGBColor(0xCC, 0xFF, 0xCC)
        r2.italic = True

    for row_idx, (row_ja, row_en) in enumerate(zip(rows_ja, rows_en)):
        row = table.rows[row_idx + 1]
        if row_idx % 2 == 1:
            for cell in row.cells:
                set_cell_bg(cell, TABLE_HEAD)
        for i, (ja, en) in enumerate(zip(row_ja, row_en)):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            r = p.add_run(f"{ja}\n")
            r.font.size = Pt(9.5)
            r2 = p.add_run(en)
            r2.font.size = Pt(8)
            r2.font.color.rgb = STONE
            r2.italic = True

    doc.add_paragraph()


def note(doc, ja, en):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(f"⚠  {ja}")
    r.font.size = Pt(10)
    r.font.color.rgb = AMBER
    r.bold = True
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent  = Cm(0.5)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(10)
    r2 = p2.add_run(f"    {en}")
    r2.font.size = Pt(9)
    r2.font.color.rgb = AMBER
    r2.italic = True


# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("業務委託請求書確認・保管ツール")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = GREEN_DARK

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("SDC Contractor Invoice Verification & Filing Tool")
r2.font.size = Pt(13); r2.font.color.rgb = STONE; r2.italic = True

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(20)
r3 = p3.add_run("ユーザーガイド  /  User Guide")
r3.bold = True; r3.font.size = Pt(15); r3.font.color.rgb = GREEN_MID

note(doc,
     "このツールは請求書の確認・保管を行います。支払いは実行しません。",
     "This tool verifies and files invoices only. It does NOT execute payments.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════

section_title(doc, "1.  ツール概要", "1.  Overview")

p = doc.add_paragraph(style="List Bullet")
r = p.add_run("毎月の業務委託請求書を自動でチェック・Google Driveへ保管するツールです。")
r.font.size = Pt(10.5)
p2 = doc.add_paragraph(style="List Bullet")
r2 = p2.add_run("Automatically checks contractor invoices and files them to Google Drive each month.")
r2.font.size = Pt(9); r2.font.color.rgb = STONE; r2.italic = True
doc.add_paragraph()

ref_table(doc,
    ["役割", "担当業務"],
    ["Role", "Responsibilities"],
    [
        ["管理者 / SDC担当者", "CSVアップロード・バリデーション・承認・Drive保管"],
        ["業務委託メンバー",   "Teamsリマインダーを受信（未提出・期日超過の場合）"],
    ],
    [
        ["Admin / SDC Staff",    "Upload CSV, validate, approve, file to Drive"],
        ["Contractor Members",   "Receive Teams reminders when invoice is missing or overdue"],
    ],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — MONTHLY WORKFLOW
# ══════════════════════════════════════════════════════════════════════════════

section_title(doc, "2.  月次ワークフロー", "2.  Monthly Workflow")

step_block(doc, 1,
    "サインイン",
    "Sign In",
    ["アプリのURLにアクセスする",
     "SDCアカウントのメールアドレスとパスワードを入力してサインインする"],
    ["Go to the app URL",
     "Enter your SDC account email and password, then sign in"],
)

step_block(doc, 2,
    "請求書CSVをアップロードする",
    "Upload the Invoice CSV",
    ["左サイドバーの「請求書一覧」をクリックする",
     "右上の「CSVアップロード」ボタンをクリックする",
     "提出システムからエクスポートしたCSV / Excelファイルを選択する",
     "カラムマッピングが正しいか確認する（自動マッピングされます）"],
    ["Click 'Invoices' in the left sidebar",
     "Click 'Upload Excel' (top right)",
     "Select the CSV / Excel file exported from the submission system",
     "Verify the column mapping is correct (auto-mapped)"],
)

step_block(doc, 3,
    "ダッシュボードで読み込み・バリデーションを実行する",
    "Load & Validate on the Dashboard",
    ["左サイドバーの「ダッシュボード」をクリックする",
     "右上の月セレクターで対象月を選択する",
     "「請求書を読み込む」ボタンをクリックする",
     "「バリデーション実行」ボタンをクリックする（自動チェック開始）"],
    ["Click 'Dashboard' in the sidebar",
     "Select the target month using the month selector (top right)",
     "Click 'Load Invoices'",
     "Click 'Run Validation' — automatic checks begin"],
)

# Validation checks table
p = doc.add_paragraph()
r = p.add_run("　バリデーション チェック項目  /  Validation Checks")
r.bold = True; r.font.size = Pt(10); r.font.color.rgb = GREEN_MID

ref_table(doc,
    ["チェック項目 / Check", "内容 / Meaning"],
    ["Check", "Meaning"],
    [
        ["PDFアクセス可能",   "添付リンクが正常に開けるか"],
        ["請求書日付あり",    "PDFから日付が抽出できたか"],
        ["消費税含む",        "消費税（10%）が確認できるか"],
        ["小計・合計あり",   "PDFから金額が読み取れたか"],
        ["金額シート一致",   "PDF合計と提出された金額が一致するか"],
        ["重複なし",          "同じファイルが過去に保管されていないか"],
    ],
    [
        ["PDF Accessible",         "Attachment link opens correctly"],
        ["Invoice Date Found",     "A date was extracted from the PDF"],
        ["Tax Included",           "Consumption tax (10%) is present"],
        ["Subtotal / Total Found", "Amounts were read from the PDF"],
        ["Amount Matches Sheet",   "PDF total matches submitted amount"],
        ["No Duplicate",           "File hasn't been filed before"],
    ],
)

step_block(doc, 4,
    "ステータスカードを確認する",
    "Review the Stats Cards",
    ["各カードをクリックすると該当する請求書一覧にフィルタリングされる",
     "「処理可能」→ そのまま保管可",
     "「要確認」→ 手動確認が必要（Step 5へ）",
     "「添付ファイルなし」→ 業務委託先に連絡",
     "「エラー」→ 処理ログを確認（Step 7へ）"],
    ["Click any card to filter the invoice list",
     "'Ready' → Can be filed immediately",
     "'Review Required' → Manual review needed (go to Step 5)",
     "'Missing Attachment' → Contact the contractor",
     "'Errors' → Check the Logs page (go to Step 7)"],
)

step_block(doc, 5,
    "「要確認」請求書を処理する",
    "Handle 'Review Required' Invoices",
    ["「要確認」カードをクリックしてフィルタリングする",
     "各行の「詳細」ボタンをクリックして右側パネルを開く",
     "失敗したチェック内容と理由を確認する",
     "内容に問題がなければ「承認」ボタンをクリックする",
     "承認された請求書はStep 6で「処理可能」として保管される"],
    ["Click 'Review Required' to filter the list",
     "Click 'View' on any row to open the detail panel (right side)",
     "Check which validation check failed and why",
     "If the invoice is acceptable despite the issue, click 'Approve'",
     "Approved invoices will be filed as Ready in Step 6"],
)

step_block(doc, 6,
    "処理可能ファイルをGoogle Driveへ保管する",
    "Save Ready Files to Google Drive",
    ["ダッシュボードの「保存可能ファイルを保存」ボタンをクリックする",
     "「処理可能」および「承認済み」の全請求書が自動でDriveへ保管される",
     "完了後、緑のバナーで保管件数を確認する"],
    ["Click 'Save Ready Files' on the Dashboard",
     "All Ready and Approved invoices are automatically filed to the correct Google Drive folder",
     "A green banner confirms how many files were saved"],
)

step_block(doc, 7,
    "処理ログを確認する",
    "Verify in the Logs",
    ["左サイドバーの「処理ログ」をクリックする",
     "左パネルで最新の実行を選択する",
     "全行が「成功」であることを確認する",
     "「エラー」や「警告」がある行は内容を確認・対応する"],
    ["Click 'Logs' in the sidebar",
     "Select the latest run from the left panel",
     "Confirm all rows show OK",
     "Investigate any ERROR or WARNING rows"],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — TEAMS REMINDERS
# ══════════════════════════════════════════════════════════════════════════════

section_title(doc, "3.  自動リマインダー（Teams）", "3.  Automatic Reminders (Teams)")

p = doc.add_paragraph(style="List Bullet")
r = p.add_run("毎朝 09:00（JST）に自動送信されます。手動送信はダッシュボードの「リマインダーを送信」から可能です。")
r.font.size = Pt(10.5)
p2 = doc.add_paragraph(style="List Bullet")
r2 = p2.add_run("Auto-sent every morning at 09:00 JST. Manual send available from Dashboard → Send Reminders.")
r2.font.size = Pt(9); r2.font.color.rgb = STONE; r2.italic = True
doc.add_paragraph()

ref_table(doc,
    ["種別", "送信条件"],
    ["Type", "When it triggers"],
    [
        ["未提出リマインダー",        "当月に請求書を提出していない業務委託先がいる"],
        ["不備・承認待ちリマインダー", "「要確認」状態が設定日数以上続いている"],
        ["期日接近アラート",           "支払期日が閾値（デフォルト：5日）以内"],
        ["期日超過アラート",           "支払期日をすでに過ぎている"],
    ],
    [
        ["Missing Invoice",      "A contractor hasn't submitted for the current month"],
        ["Stale Review",         "Invoice stuck in Review Required for too many days"],
        ["Due Date Approaching", "Payment due date within threshold (default 5 days)"],
        ["Overdue",              "Payment due date has already passed"],
    ],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — STATUS REFERENCE
# ══════════════════════════════════════════════════════════════════════════════

section_title(doc, "4.  ステータスコード 早見表", "4.  Status Code Quick Reference")

ref_table(doc,
    ["ステータス", "色", "意味・対応"],
    ["Status", "Colour", "Meaning & Action"],
    [
        ["処理可能",    "緑",  "全チェック通過 → 保管可"],
        ["要確認",      "橙",  "チェック失敗 → Step 5で手動確認"],
        ["添付なし",    "赤",  "PDF添付なし → 業務委託先に連絡"],
        ["金額不一致",  "赤",  "PDF合計 ≠ シート金額 → 確認・承認"],
        ["処理済み",    "灰",  "過去runで保管済み → 対応不要"],
        ["保存済み",    "青",  "Drive保管完了 → 完了"],
        ["保存エラー",  "赤",  "保管失敗 → 処理ログを確認"],
    ],
    [
        ["Ready",              "Green",  "All checks passed → safe to file"],
        ["Review Required",   "Amber",  "Check failed → manual review in Step 5"],
        ["Missing Attachment", "Red",   "No PDF → contact contractor"],
        ["Amount Mismatch",    "Red",   "PDF ≠ sheet → review and approve"],
        ["Already Processed",  "Grey",  "Filed previously → no action needed"],
        ["Saved",              "Blue",  "Filed to Drive → done"],
        ["Save Error",         "Red",   "Filing failed → check Logs"],
    ],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — TIPS
# ══════════════════════════════════════════════════════════════════════════════

section_title(doc, "5.  利用上のポイント", "5.  Tips")

tips = [
    ("対象月を必ず確認してから読み込んでください（デフォルトは当月）",
     "Always confirm the correct month before loading — defaults to current month"),
    ("「承認」はチェック内容を手動確認した場合のみ使用してください",
     "Use Approve only when you have manually confirmed the invoice is correct"),
    ("毎回の処理後に処理ログを確認してエラーがないか確かめてください",
     "Check the Logs page after every run to confirm no errors occurred"),
    ("このツールは支払いを実行しません。Drive保管が最終ステップです",
     "This tool does not execute payments — filing to Drive is the final step"),
]

for ja, en in tips:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(ja)
    r.font.size = Pt(10.5)
    p2 = doc.add_paragraph(style="List Bullet")
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(8)
    r2 = p2.add_run(en)
    r2.font.size = Pt(9); r2.font.color.rgb = STONE; r2.italic = True

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SDC 業務委託請求書確認・保管ツール — 社内資料  /  Internal Documentation — June 2026")
r.font.size = Pt(8); r.font.color.rgb = STONE

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save("SDC_Invoice_Tool_User_Guide.docx")
print("Saved: SDC_Invoice_Tool_User_Guide.docx")
